import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import RGBColor, Inches, Pt
from pptx import Presentation
from pptx.util import Inches as PPTInches, Pt as PPTPt
from pptx.dml.color import RGBColor as PPTColor
from pptx.enum.shapes import MSO_SHAPE
from io import BytesIO
from datetime import datetime

# ════════════════════════════════════════════════════════════
# PAGE CONFIG
# ════════════════════════════════════════════════════════════
st.set_page_config(
    page_title="PDD Automation Suite",
    page_icon="📋",
    layout="wide"
)

# ════════════════════════════════════════════════════════════
# SESSION STATE - Persistent across all tabs
# ════════════════════════════════════════════════════════════
DEFAULTS = {
    "gids_number": "",
    "process_name": "",
    "business_unit": "",
    "process_owner": "",
    "processing_steps": [],
    "business_rules": [],
    "aht_minutes": 0,
    "daily_volume": 0,
    "monthly_volume": 0,
    "peak_periods": "",
    "target_sla": "",
    "out_of_scope": "",
    "applications": [],
    "as_is_flow": "",
    "to_be_flow": "",
}

for k, v in DEFAULTS.items():
    if k not in st.session_state:
        st.session_state[k] = v

# ════════════════════════════════════════════════════════════
# SIDEBAR
# ════════════════════════════════════════════════════════════
with st.sidebar:
    st.title("⚙️ Controls")
    st.info("✅ Default IA-compliant template active")
    st.divider()
    st.markdown("### 🔒 Privacy")
    if st.button("Wipe All Session Data", use_container_width=True, type="primary"):
        for k in list(st.session_state.keys()):
            del st.session_state[k]
        st.rerun()
    st.caption("🛡️ Data lives in your browser session only. Nothing is saved on the server.")
    st.divider()
    st.caption("Built with ❤️ using Streamlit")

# ════════════════════════════════════════════════════════════
# HEADER
# ════════════════════════════════════════════════════════════
st.title("📋 PDD / HLD / LLD Automation Suite")
st.caption("Fill any tab in any order — your data is preserved automatically across all tabs.")

# ════════════════════════════════════════════════════════════
# TABS (STEPPER)
# ════════════════════════════════════════════════════════════
tabs = st.tabs([
    "1️⃣ Metadata",
    "2️⃣ Steps",
    "3️⃣ Rules",
    "4️⃣ Metrics",
    "5️⃣ Apps",
    "6️⃣ As-Is / To-Be",
    "🚀 Generate"
])

# ─────────── TAB 1: METADATA ───────────
with tabs[0]:
    st.subheader("Process Metadata")
    c1, c2 = st.columns(2)
    with c1:
        st.session_state["gids_number"] = st.text_input(
            "GIDS Number *", st.session_state["gids_number"]
        )
        st.session_state["process_name"] = st.text_input(
            "Process Name *", st.session_state["process_name"]
        )
    with c2:
        st.session_state["business_unit"] = st.text_input(
            "Business Unit *", st.session_state["business_unit"]
        )
        st.session_state["process_owner"] = st.text_input(
            "Process Owner", st.session_state["process_owner"]
        )

# ─────────── TAB 2: PROCESSING STEPS ───────────
with tabs[1]:
    st.subheader("Processing Steps (As-Is Execution)")
    st.caption("Add each step the user currently performs manually.")
    if st.session_state["processing_steps"]:
        df_steps = pd.DataFrame(st.session_state["processing_steps"])
    else:
        df_steps = pd.DataFrame(columns=["Step", "Action", "System", "Input", "Expected Output"])
    edited_steps = st.data_editor(
        df_steps,
        num_rows="dynamic",
        use_container_width=True,
        key="steps_editor"
    )
    st.session_state["processing_steps"] = edited_steps.to_dict("records")

# ─────────── TAB 3: BUSINESS RULES ───────────
with tabs[2]:
    st.subheader("Business Rules & Decision Logic")
    st.caption("Capture validations, conditions, and decision matrix logic.")
    if st.session_state["business_rules"]:
        df_rules = pd.DataFrame(st.session_state["business_rules"])
    else:
        df_rules = pd.DataFrame(columns=["Rule ID", "Condition", "Action / Outcome"])
    edited_rules = st.data_editor(
        df_rules,
        num_rows="dynamic",
        use_container_width=True,
        key="rules_editor"
    )
    st.session_state["business_rules"] = edited_rules.to_dict("records")

# ─────────── TAB 4: METRICS & SCOPE ───────────
with tabs[3]:
    st.subheader("Metrics & Scope")
    c1, c2 = st.columns(2)
    with c1:
        st.session_state["aht_minutes"] = st.number_input(
            "Average Handling Time (minutes)", 0, 1000,
            int(st.session_state["aht_minutes"])
        )
        st.session_state["daily_volume"] = st.number_input(
            "Daily Transaction Volume", 0, 1000000,
            int(st.session_state["daily_volume"])
        )
        st.session_state["monthly_volume"] = st.number_input(
            "Monthly Transaction Volume", 0, 10000000,
            int(st.session_state["monthly_volume"])
        )
    with c2:
        st.session_state["target_sla"] = st.text_input(
            "Target SLA (e.g., 24 hrs)", st.session_state["target_sla"]
        )
        st.session_state["peak_periods"] = st.text_area(
            "Peak Scaling Periods", st.session_state["peak_periods"], height=100
        )
    st.session_state["out_of_scope"] = st.text_area(
        "Out-of-Scope Scenarios", st.session_state["out_of_scope"], height=120
    )

# ─────────── TAB 5: APPS & ENVIRONMENT ───────────
with tabs[4]:
    st.subheader("Applications & Environment")
    st.caption("List all target systems involved in this process.")
    if st.session_state["applications"]:
        df_apps = pd.DataFrame(st.session_state["applications"])
    else:
        df_apps = pd.DataFrame(columns=["Application", "Type", "Credentials Profile"])
    edited_apps = st.data_editor(
        df_apps,
        num_rows="dynamic",
        use_container_width=True,
        key="apps_editor"
    )
    st.session_state["applications"] = edited_apps.to_dict("records")

# ─────────── TAB 6: AS-IS / TO-BE ───────────
with tabs[5]:
    st.subheader("As-Is vs To-Be Workflow")
    c1, c2 = st.columns(2)
    with c1:
        st.session_state["as_is_flow"] = st.text_area(
            "As-Is (Manual Process)",
            st.session_state["as_is_flow"],
            height=300
        )
    with c2:
        st.session_state["to_be_flow"] = st.text_area(
            "To-Be (Automated Process)",
            st.session_state["to_be_flow"],
            height=300
        )

# ════════════════════════════════════════════════════════════
# GENERATOR FUNCTIONS
# ════════════════════════════════════════════════════════════

def add_table_from_data(doc, data_list, fallback_cols):
    """Helper: add a styled table from list-of-dicts."""
    if not data_list:
        doc.add_paragraph("No data provided.")
        return
    cols = list(data_list[0].keys())
    table = doc.add_table(rows=1, cols=len(cols))
    try:
        table.style = "Light Grid Accent 1"
    except KeyError:
        table.style = "Table Grid"
    for i, c in enumerate(cols):
        cell = table.rows[0].cells[i]
        cell.text = str(c)
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
    for row in data_list:
        cells = table.add_row().cells
        for i, c in enumerate(cols):
            cells[i].text = str(row.get(c, ""))

def build_pdd():
    """Build the PDD .docx document."""
    doc = Document()

    # Title page
    title = doc.add_heading("Process Definition Document", 0)
    title.alignment = 1
    doc.add_paragraph(f"Process: {st.session_state['process_name'] or 'N/A'}").alignment = 1
    doc.add_paragraph(f"GIDS: {st.session_state['gids_number'] or 'N/A'}").alignment = 1
    doc.add_paragraph(f"Generated: {datetime.now():%Y-%m-%d %H:%M}").alignment = 1
    doc.add_page_break()

    # 1. Metadata
    doc.add_heading("1. Process Metadata", 1)
    meta = [
        ("GIDS Number", st.session_state["gids_number"]),
        ("Process Name", st.session_state["process_name"]),
        ("Business Unit", st.session_state["business_unit"]),
        ("Process Owner", st.session_state["process_owner"]),
    ]
    t = doc.add_table(rows=len(meta), cols=2)
    try:
        t.style = "Light Grid Accent 1"
    except KeyError:
        t.style = "Table Grid"
    for i, (k, v) in enumerate(meta):
        t.cell(i, 0).text = k
        t.cell(i, 1).text = str(v) if v else "N/A"

    # 2. As-Is / To-Be
    doc.add_heading("2. As-Is vs To-Be Workflow", 1)
    doc.add_heading("As-Is (Manual)", 2)
    doc.add_paragraph(st.session_state["as_is_flow"] or "Not provided")
    doc.add_heading("To-Be (Automated)", 2)
    doc.add_paragraph(st.session_state["to_be_flow"] or "Not provided")

    # 3. Processing Steps
    doc.add_heading("3. Processing Steps", 1)
    add_table_from_data(
        doc, st.session_state["processing_steps"],
        ["Step", "Action", "System", "Input", "Expected Output"]
    )

    # 4. Business Rules
    doc.add_heading("4. Business Rules & Decision Logic", 1)
    add_table_from_data(
        doc, st.session_state["business_rules"],
        ["Rule 
